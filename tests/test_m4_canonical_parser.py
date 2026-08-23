from __future__ import annotations

import hashlib
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document
from PIL import Image
from reportlab.pdfgen import canvas

from astra_indexator.acquisition import AcquiredSource
from astra_indexator.parser import (
    DocumentParserService,
    ElementType,
    FileTypeHandlerRegistry,
    ImageDocumentHandler,
    MarkdownDocumentHandler,
    ParseContext,
    ParserError,
    PdfDocumentHandler,
    TextDocumentHandler,
    DocxDocumentHandler,
    default_registry,
)
from astra_indexator.storage import StorageRef


def _source(path: Path, fmt: str) -> AcquiredSource:
    payload = path.read_bytes()
    return AcquiredSource(
        source_ref=StorageRef("documents", path.name),
        local_path=path,
        original_file_name=path.name,
        detected_format=fmt,
        detected_content_type="application/octet-stream",
        size_bytes=len(payload),
        sha256=hashlib.sha256(payload).hexdigest(),
        etag=None,
        version_id=None,
        validation_profile="default-v1",
        warnings=(),
        acquired_at=datetime.now(timezone.utc),
    )


def _ctx(source: AcquiredSource) -> ParseContext:
    return ParseContext(uuid4(), uuid4(), uuid4(), 1, source.sha256)


def test_registry_rejects_duplicate_format_handlers() -> None:
    with pytest.raises(ValueError):
        FileTypeHandlerRegistry([TextDocumentHandler(), TextDocumentHandler()])


def test_registry_does_not_fallback_from_detected_format() -> None:
    registry = FileTypeHandlerRegistry([TextDocumentHandler()])
    with pytest.raises(ParserError) as exc:
        registry.resolve("PDF")
    assert exc.value.code == "PARSER_UNSUPPORTED_FORMAT"


def test_txt_is_structured_and_deterministic(tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"
    path.write_text("Первый абзац.\n\nЕкінші абзац.\n\nThird paragraph.", encoding="utf-8")
    source = _source(path, "TXT")
    ctx = _ctx(source)
    parser = DocumentParserService(default_registry())
    first = parser.parse(source, ctx)
    second = parser.parse(source, ctx)
    assert [e.type for e in first.elements] == [ElementType.PARAGRAPH] * 3
    assert [e.element_id for e in first.elements] == [e.element_id for e in second.elements]
    assert first.quality.native_text_chars > 0


def test_markdown_preserves_headings_lists_and_code(tmp_path: Path) -> None:
    path = tmp_path / "sample.md"
    path.write_text("# Title\n\nParagraph.\n\n- one\n- two\n\n```python\nprint('x')\n```\n", encoding="utf-8")
    source = _source(path, "MARKDOWN")
    result = DocumentParserService(default_registry()).parse(source, _ctx(source))
    kinds = [e.type for e in result.elements]
    assert ElementType.HEADING in kinds
    assert kinds.count(ElementType.LIST_ITEM) == 2
    assert ElementType.CODE_BLOCK in kinds


def test_docx_preserves_heading_paragraph_and_table(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Architecture", level=1)
    doc.add_paragraph("Canonical paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.save(path)
    source = _source(path, "DOCX")
    result = DocumentParserService(default_registry()).parse(source, _ctx(source))
    assert result.elements[0].type == ElementType.HEADING
    assert any(e.type == ElementType.TABLE and e.metadata["rowCount"] == 2 for e in result.elements)
    assert all(e.geometry is None for e in result.elements)


def test_image_becomes_first_class_element_and_ocr_candidate(tmp_path: Path) -> None:
    path = tmp_path / "scan.png"
    Image.new("RGB", (320, 200), "white").save(path)
    source = _source(path, "PNG")
    result = DocumentParserService(default_registry()).parse(source, _ctx(source))
    assert len(result.elements) == 1
    assert result.elements[0].type == ElementType.IMAGE
    assert result.elements[0].geometry.page_width == 320
    assert len(result.ocr_candidates) == 1
    assert result.quality.status.value == "OCR_REQUIRED"


def test_pdf_native_text_has_page_provenance_and_page_break(tmp_path: Path) -> None:
    path = tmp_path / "native.pdf"
    c = canvas.Canvas(str(path))
    c.drawString(72, 760, "Native PDF heading")
    c.drawString(72, 730, "First paragraph line")
    c.showPage()
    c.drawString(72, 760, "Second page")
    c.save()
    source = _source(path, "PDF")
    result = DocumentParserService(default_registry()).parse(source, _ctx(source))
    text_elements = [e for e in result.elements if e.text]
    assert text_elements
    assert {e.geometry.page_number for e in text_elements} == {1, 2}
    assert sum(1 for e in result.elements if e.type == ElementType.PAGE_BREAK) == 2
    assert len(result.quality.page_modes) == 2


def test_parser_rejects_context_hash_mismatch(tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"
    path.write_text("text", encoding="utf-8")
    source = _source(path, "TXT")
    ctx = ParseContext(uuid4(), uuid4(), uuid4(), 1, "0" * 64)
    with pytest.raises(ParserError) as exc:
        DocumentParserService(default_registry()).parse(source, ctx)
    assert exc.value.code == "PARSER_SOURCE_HASH_MISMATCH"
