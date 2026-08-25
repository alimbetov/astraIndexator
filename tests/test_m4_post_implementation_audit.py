from __future__ import annotations

import hashlib
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.shared import Inches
from PIL import Image

from astra_indexator.acquisition import AcquiredSource
from astra_indexator.parser import (
    DocumentParserService,
    ElementType,
    ParseContext,
    QualityStatus,
    default_registry,
)
from astra_indexator.storage import StorageRef


def _source(path: Path, fmt: str) -> AcquiredSource:
    payload = path.read_bytes()
    return AcquiredSource(
        source_ref=StorageRef.parse(f"seaweed://documents/{path.name}"),
        local_path=path,
        original_file_name=path.name,
        detected_format=fmt,
        detected_content_type="application/octet-stream",
        size_bytes=len(payload),
        sha256=hashlib.sha256(payload).hexdigest(),
        etag=None,
        version_id=None,
        validation_profile="default-v1",
        warnings=(),
        acquired_at=datetime.now(timezone.utc),
    )


def _ctx(source: AcquiredSource) -> ParseContext:
    return ParseContext(uuid4(), uuid4(), uuid4(), 1, source.sha256)


def test_docx_embedded_image_keeps_body_order_without_forcing_whole_document_ocr(
    tmp_path: Path,
) -> None:
    image_path = tmp_path / "embedded.png"
    Image.new("RGB", (64, 32), "white").save(image_path)

    path = tmp_path / "ordered.docx"
    document = Document()
    document.add_paragraph("Before image")
    image_paragraph = document.add_paragraph()
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("After image")
    document.save(path)

    source = _source(path, "DOCX")
    result = DocumentParserService(default_registry()).parse(source, _ctx(source))

    significant = [(element.type, element.text) for element in result.elements]
    assert significant[0] == (ElementType.PARAGRAPH, "Before image")
    assert significant[1][0] == ElementType.IMAGE
    assert significant[2] == (ElementType.PARAGRAPH, "After image")
    assert len(result.ocr_candidates) == 1
    assert result.ocr_candidates[0].scope == "EMBEDDED_IMAGE"
    assert result.quality.status == QualityStatus.GOOD


def test_standalone_image_still_requires_ocr(tmp_path: Path) -> None:
    path = tmp_path / "scan.png"
    Image.new("RGB", (100, 50), "white").save(path)
    source = _source(path, "PNG")
    result = DocumentParserService(default_registry()).parse(source, _ctx(source))
    assert result.quality.status == QualityStatus.OCR_REQUIRED


def test_rtf_partial_structure_warning_is_reflected_in_quality(tmp_path: Path) -> None:
    path = tmp_path / "sample.rtf"
    path.write_text(r"{\rtf1\ansi First paragraph.\par Second paragraph.}", encoding="latin-1")
    source = _source(path, "RTF")
    result = DocumentParserService(default_registry()).parse(source, _ctx(source))
    assert "RTF_STRUCTURE_PARTIAL" in result.quality.warnings
    assert result.quality.status == QualityStatus.PARTIAL
